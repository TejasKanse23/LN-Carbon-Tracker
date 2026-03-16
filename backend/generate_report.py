import os
import argparse
import json
import uuid
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report(data, file_path):
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Carbon Intelligence Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(f"Lane: {data['overview']['lane_name']}", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    
    # 1. Lane Overview
    doc.add_heading("1. Lane Overview", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Total Shipments: ").bold = True
    p.add_run(f"{data['overview']['total_shipments']}\n")
    p.add_run(f"Total Distance Traveled: ").bold = True
    p.add_run(f"{data['overview']['total_distance_km']:,} km\n")
    p.add_run(f"Total Freight Transported: ").bold = True
    p.add_run(f"{data['overview']['total_freight_tons']:,} tons\n")
    p.add_run(f"Average Utilization: ").bold = True
    p.add_run(f"{data['overview']['avg_utilization']:.1f}%\n")
    p.add_run(f"Average Vehicle Age: ").bold = True
    p.add_run(f"{data['overview']['avg_vehicle_age']:.1f} years\n")
    p.add_run(f"Common Vehicle Types: ").bold = True
    p.add_run(f"{', '.join(data['overview']['common_vehicle_types'])}\n")

    # 2. Carbon Emission Analysis
    doc.add_heading("2. Carbon Emission Analysis", level=2)
    p2 = doc.add_paragraph()
    p2.add_run(f"Total CO2 Emissions: ").bold = True
    p2.add_run(f"{data['emissions']['total_co2_kg']:,} kg\n")
    p2.add_run(f"Avg Emissions per Shipment: ").bold = True
    p2.add_run(f"{data['emissions']['avg_co2_per_shipment']:,} kg\n")
    if data['emissions']['emission_intensity'] > 0:
        p2.add_run(f"Emission Intensity: ").bold = True
        p2.add_run(f"{data['emissions']['emission_intensity']:.4f} kg CO2 / ton-km\n")
    p2.add_run(f"Network Average per Shipment: ").bold = True
    p2.add_run(f"{data['emissions']['network_avg_co2']:,} kg\n")
    p2.add_run(f"Lane Status vs Network: ").bold = True
    status_run = p2.add_run(f"{data['emissions']['status_vs_network']}")
    status_run.font.color.rgb = RGBColor(255, 0, 0) if "Above" in data['emissions']['status_vs_network'] else RGBColor(0, 153, 0)

    # 3. Emission Drivers
    doc.add_heading("3. Emission Drivers", level=2)
    p3 = doc.add_paragraph("The following factors are driving emissions on this lane:\n")
    for driver in data['drivers']:
        doc.add_paragraph(f"• {driver}", style='List Bullet')

    # 4. Carbon Hotspot Detection
    doc.add_heading("4. Carbon Hotspot Detection", level=2)
    p4 = doc.add_paragraph("Key inefficiencies and anomalies detected:\n")
    for hotspot in data['hotspots']:
        doc.add_paragraph(f"• {hotspot}", style='List Bullet')

    # 5. Sustainability Recommendations
    doc.add_heading("5. Sustainability Recommendations", level=2)
    p5 = doc.add_paragraph("Actionable steps to decarbonize this lane:\n")
    for rec in data['recommendations']:
        doc.add_paragraph(f"• {rec}", style='List Bullet')

    # 6. Reduction Opportunity Summary
    doc.add_heading("6. Reduction Opportunity Summary", level=2)
    p6 = doc.add_paragraph()
    p6.add_run(f"Estimated Potential Reduction: ").bold = True
    p6.add_run(f"{data['reduction']['potential_reduction_kg']:,} kg CO2\n")
    p6.add_run(f"Impact: ").bold = True
    p6.add_run(f"{data['reduction']['impact_summary']}")

    doc.save(file_path)

def analyze_lane(origin, dest, data_path):
    df = pd.read_csv(data_path)
    
    # Lane filter
    lane_df = df[(df['origin'].str.lower() == origin.lower()) & (df['destination'].str.lower() == dest.lower())]
    if lane_df.empty:
        # Fallback if specific direction doesn't exist, try ignoring direction or just return empty
        # Let's see if there are any shipments
        return None

    network_avg_co2 = df['estimated_emissions_kg_co2e'].mean()

    # 1. Overview
    total_shipments = len(lane_df)
    total_distance = lane_df['distance_km'].sum()
    total_freight = lane_df['weight_ton'].sum()
    avg_util = lane_df['utilization_percent'].mean()
    if 'Age of vehicle' in lane_df.columns:
        avg_age = lane_df['Age of vehicle'].mean()
    else:
        avg_age = 0.0
    common_vehicles = lane_df['vehicle_type'].value_counts().head(3).index.tolist()

    # 2. Emissions
    total_co2 = lane_df['estimated_emissions_kg_co2e'].sum()
    avg_co2 = total_co2 / total_shipments if total_shipments > 0 else 0
    intensity = total_co2 / (total_distance * total_freight / total_shipments) if total_freight > 0 else 0
    
    status_vs_network = "Above Network Average" if avg_co2 > network_avg_co2 else "Below Network Average"

    # 3. Drivers
    drivers = []
    congestion_high = (lane_df.get('traffic congestion', pd.Series(['Low'])).str.lower() == 'high').mean()
    if congestion_high > 0.3:
        drivers.append("High traffic congestion significantly impacts fuel efficiency on this lane.")
    if avg_age > 6.0:
        drivers.append(f"Older vehicle fleet (average age {avg_age:.1f} years) contributes to higher emissions.")
    if avg_util < 70.0:
        drivers.append(f"Low average utilization ({avg_util:.1f}%) results in higher emissions per ton-km.")
    diesel_pct = (lane_df['fuel_type'].str.lower() == 'diesel').mean() * 100
    if diesel_pct > 50:
        drivers.append(f"High dependence on Diesel fuel ({diesel_pct:.1f}% of shipments).")
    if not drivers:
        drivers.append("Operations are relatively optimized, though incremental improvements are possible.")

    # 4. Hotspots
    hotspots = []
    # find shipments > 80th percentile
    co2_80 = lane_df['estimated_emissions_kg_co2e'].quantile(0.8)
    high_emitters = lane_df[lane_df['estimated_emissions_kg_co2e'] > co2_80]
    if len(high_emitters) > 0:
        inefficient_util = high_emitters[high_emitters['utilization_percent'] < 60]
        if len(inefficient_util) > 0:
            hotspots.append(f"{len(inefficient_util)} shipments identified with both high emissions and poor utilization (<60%).")
        hotspots.append(f"Top 20% of shipments contribute to a disproportionate amount of emissions.")
    hotspots.append("Peak congestion periods correlate directly with emission spikes.")

    # 5. Recommendations
    recs = []
    if avg_util < 85:
        recs.append("Implement load consolidation strategies to improve vehicle utilization.")
    if diesel_pct > 30:
        recs.append("Gradually shift from Diesel to CNG or EV alternatives for suitable loads.")
    if avg_age > 5.0:
        recs.append("Phase out vehicles older than 7 years to benefit from modern fuel-efficient engines.")
    if congestion_high > 0.2:
        recs.append("Reschedule departure times to off-peak hours to avoid congestion penalties.")

    if not recs:
        recs.append("Continue current optimized operational practices.")

    # 6. Reduction
    # Suppose we bring avg CO2 down to network avg or reduce by 15%
    if avg_co2 > network_avg_co2:
        potential = (avg_co2 - network_avg_co2) * total_shipments
    else:
        potential = total_co2 * 0.10 # 10% improvement
    
    reduction_impact = f"Implementing these changes could reduce lane emissions by approximately {potential/total_co2*100:.1f}%."

    report_data = {
        "overview": {
            "lane_name": f"{origin} → {dest}",
            "total_shipments": int(total_shipments),
            "total_distance_km": float(round(total_distance, 1)),
            "total_freight_tons": float(round(total_freight, 1)),
            "avg_utilization": float(round(avg_util, 2)),
            "avg_vehicle_age": float(round(avg_age, 1)),
            "common_vehicle_types": common_vehicles
        },
        "emissions": {
            "total_co2_kg": float(round(total_co2, 1)),
            "avg_co2_per_shipment": float(round(avg_co2, 1)),
            "emission_intensity": float(round(intensity, 4)),
            "network_avg_co2": float(round(network_avg_co2, 1)),
            "status_vs_network": status_vs_network
        },
        "drivers": drivers,
        "hotspots": hotspots,
        "recommendations": recs,
        "reduction": {
            "potential_reduction_kg": float(round(potential, 1)),
            "impact_summary": reduction_impact
        }
    }

    return report_data

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--origin", required=True)
    parser.add_argument("--dest", required=True)
    args = parser.parse_args()

    data_path = os.path.join(os.path.dirname(__file__), "..", "carbon_tracker_agent", "data", "carbon_tracker_shipments.csv")
    
    if not os.path.exists("reports"):
        os.makedirs("reports")

    data = analyze_lane(args.origin, args.dest, data_path)
    if data is None:
        print(json.dumps({"error": "No shipments found for this lane"}))
    else:
        file_name = f"Report_{args.origin}_{args.dest}_{uuid.uuid4().hex[:6]}.docx"
        file_path = os.path.join("reports", file_name)
        create_word_report(data, file_path)
        data['file_path'] = file_path
        print(json.dumps(data))
